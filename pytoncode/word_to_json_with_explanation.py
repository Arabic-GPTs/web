
import docx
import json
import re
from pathlib import Path

"""
هذا السكريبت يقوم باستخراج المحتوى من ملف Word (docx) وتحويله إلى هيكل JSON منظم.
تعتمد آلية العمل على تحليل الأنماط النصية داخل ملف Word لتحديد العناوين الرئيسية، العناوين الفرعية، عناصر القائمة، وتفاصيل كل عنصر.

الخطوات المتبعة:

1.  **تحليل هيكل ملف Word:**
    *   تم فهم كيفية تنظيم المحتوى داخل ملف Word المرفق، حيث توجد أنماط متكررة للعناوين الرئيسية (`العنوان الرئيسي:`), العناوين الفرعية (`العنوان الفرعي:`), وعناوين العناصر الفردية (`#`).
    *   تم تحديد التفاصيل لكل عنصر تبدأ بالرمز (`@`) متبوعًا بالمفتاح والقيمة، بالإضافة إلى الروابط.

2.  **تثبيت المكتبات اللازمة:**
    *   تم استخدام مكتبة `python-docx` لاستخراج النصوص من ملفات `.docx`. يمكن تثبيتها باستخدام `pip3 install python-docx`.

3.  **آلية عمل السكريبت (`extract_content_from_docx`):**
    *   **قراءة المستند:** يستخدم `docx.Document(docx_file_path)` لفتح وقراءة ملف Word.
    *   **التكرار على الفقرات:** يقوم السكريبت بالمرور على كل فقرة في المستند.
    *   **تحديد الهيكل:**
        *   **العنوان الرئيسي:** إذا بدأت الفقرة بـ `العنوان الرئيسي:`، يتم اعتبارها عنوانًا رئيسيًا جديدًا ويتم إنشاء مفتاح جديد في قاموس JSON الرئيسي.
        *   **العنوان الفرعي:** إذا بدأت الفقرة بـ `العنوان الفرعي:`، يتم اعتبارها عنوانًا فرعيًا جديدًا ضمن العنوان الرئيسي الحالي. تم التعامل مع حالة خاصة حيث قد يأتي العنوان الفرعي وعنوان العنصر الأول في نفس السطر (مثال: `العنوان الفرعي: نماذج الابتكار#اقتراح عنوان وفكرة بحث`)، حيث يتم فصلهما بشكل صحيح.
        *   **عناوين العناصر:** إذا بدأت الفقرة بـ `#`، يتم اعتبارها عنوان عنصر جديد ضمن العنوان الفرعي الحالي، ويتم إنشاء كائن جديد في قائمة العناصر.
        *   **تفاصيل العناصر:** إذا بدأت الفقرة بـ `@`، يتم اعتبارها تفصيلاً للعنصر الحالي (مثل نبذة، حدود، مثال، روابط). يتم تجميع الأسطر المتتالية التي لا تبدأ بأي من العلامات المذكورة كجزء من قيمة التفصيل الحالي (للسماح بالفقرات متعددة الأسطر).
    *   **معالجة الروابط:** يتم استخراج الروابط بشكل خاص، حيث يتم تقسيمها إذا كانت متعددة في سطر واحد وإضافتها إلى قائمة الروابط. كما يتم التعامل مع حقول `نموذج 4o` و `نموذج 5` كروابط إذا كانت تحتوي على `https://`.
    *   **تنظيف القيم:** يتم تطبيق دالة `clean_value` لإزالة المسافات البيضاء الزائدة وعلامات الاقتباس من بداية ونهاية النصوص.
    *   **حفظ التفاصيل:** يتم استخدام دالة `save_current_detail()` لحفظ التفاصيل المجمعة عند الانتقال إلى عنوان رئيسي أو فرعي أو عنصر جديد، أو عند انتهاء المستند.

4.  **تنفيذ السكريبت وإنشاء ملف JSON:**
    *   يتم تنفيذ السكريبت، الذي يقرأ ملف Word المرفق، ويعالج محتواه، ثم يحفظ الناتج في ملف JSON جديد باسم `output_from_docx.json`.

"""

def clean_value(val):
    # Remove leading/trailing whitespace and quotes
    return val.strip().strip("\'").strip("\"")

def extract_content_from_docx(docx_file_path):
    document = docx.Document(docx_file_path)
    data = {}
    current_main_title = None
    current_sub_title = None
    current_item = None
    current_detail_key = None
    current_detail_value_buffer = []

    def save_current_detail():
        nonlocal current_detail_key, current_detail_value_buffer
        if current_item and current_detail_key:
            value = " ".join([clean_value(v) for v in current_detail_value_buffer if clean_value(v)])
            if current_detail_key == "روابط":
                links = [link.strip() for link in re.split(r'\s+|\n', value) if link.strip()]
                current_item["details"]["روابط"].extend(links)
                current_item["details"]["روابط"] = list(dict.fromkeys(current_item["details"]["روابط"])) # Remove duplicates
            elif current_detail_key.startswith("نموذج"):
                if "https://" in value:
                    if "روابط" not in current_item["details"]:
                        current_item["details"]["روابط"] = []
                    current_item["details"]["روابط"].append(value)
                    current_item["details"]["روابط"] = list(dict.fromkeys(current_item["details"]["روابط"])) # Remove duplicates
                else:
                    current_item["details"][current_detail_key] = value
            else:
                current_item["details"][current_detail_key] = value
        current_detail_key = None
        current_detail_value_buffer = []

    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if not line:
            continue

        # Check for page breaks or separators that might be in the docx
        if line.startswith("--- PAGE") or line == "________________________________________":
            save_current_detail()
            continue

        if line.startswith("العنوان الرئيسي:"):
            save_current_detail()
            current_main_title = clean_value(line.replace("العنوان الرئيسي:", ""))
            data[current_main_title] = {}
            current_sub_title = None
            current_item = None
        elif line.startswith("العنوان الفرعي:"):
            save_current_detail()
            if current_main_title is not None:
                # Check if the sub-title line also contains an item title (e.g., "نماذج الابتكار#اقتراح عنوان وفكرة بحث")
                sub_title_text = line.replace("العنوان الفرعي:", "").strip()
                if '#' in sub_title_text:
                    parts = sub_title_text.split('#', 1)
                    current_sub_title = clean_value(parts[0])
                    item_title = clean_value(parts[1])
                    if current_sub_title not in data[current_main_title]:
                        data[current_main_title][current_sub_title] = []
                    current_item = {"title": item_title, "details": {}}
                    data[current_main_title][current_sub_title].append(current_item)
                    # Initialize common detail keys
                    current_item["details"]["نبذة"] = ""
                    current_item["details"]["حدود"] = ""
                    current_item["details"]["مثال"] = ""
                    current_item["details"]["روابط"] = []
                else:
                    current_sub_title = clean_value(sub_title_text)
                    data[current_main_title][current_sub_title] = []
                    current_item = None
        elif line.startswith("#"):
            save_current_detail()
            if current_main_title is not None and current_sub_title is not None:
                item_title = clean_value(line[1:])
                current_item = {"title": item_title, "details": {}}
                data[current_main_title][current_sub_title].append(current_item)
                # Initialize common detail keys
                current_item["details"]["نبذة"] = ""
                current_item["details"]["حدود"] = ""
                current_item["details"]["مثال"] = ""
                current_item["details"]["روابط"] = []
        elif line.startswith("@"):
            save_current_detail()
            if current_item is not None:
                parts = line[1:].split(": ", 1)
                key = clean_value(parts[0])
                value = clean_value(parts[1]) if len(parts) > 1 else ""
                
                current_detail_key = key
                current_detail_value_buffer.append(value)
        elif current_item is not None and current_detail_key is not None:
            # Continue appending to the current detail's value buffer
            current_detail_value_buffer.append(line)

    save_current_detail() # Save any remaining detail after the loop

    return json.dumps(data, ensure_ascii=False, indent=4)

if __name__ == "__main__":
    script_dir = Path(__file__).resolve().parent
    docx_file = script_dir / "01.docx"
    json_output = extract_content_from_docx(str(docx_file))

    output_path = script_dir / "output_from_docx.json"
    with output_path.open("w", encoding="utf-8") as f:
        f.write(json_output)
    print("تم إنشاء ملف output_from_docx.json بنجاح")

